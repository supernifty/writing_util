#!/usr/bin/env python
'''
  word doc of author list from spreadsheet

  installation:
    pip install python-docx
'''

import argparse
import csv
import logging
import sys

from docx import Document

def main(fn, name, affiliations, title, output):

  document = Document()

  p = document.add_paragraph()
  p.add_run("This file is automatically generated. Do not edit.")

  p = document.add_paragraph()
  p.add_run(title).bold = True

  logging.info('reading from %s...', fn)
  affs = {} # affnumber -> affname
  affs_seen = {} # affname -> affnumber
  additional = {}
  next_aff = 1
  authors_seen = set()
  
  p = document.add_paragraph()

  for i, row in enumerate(csv.DictReader(open(fn, mode='r', encoding='utf-8-sig'))):
    if i < 2:
      logging.debug(row)

    n = row[name].strip() # name of author

    if n == '':
      continue

    if n in authors_seen:
      logging.warn('%s is duplicated', n)
    authors_seen.add(n)

    if i > 0:
      p.add_run(', ') #.italic = True

    logging.debug('processing %s', n)

    n_affs = set() # author affiliations
    n_additional = set() # author affiliations
    for affiliation in affiliations:
      a = row[affiliation].strip().strip('.')
      if a == '':
        continue # next affiliation


      if a not in affs_seen: # it's a new affiliation
        affs[next_aff] = a
        affs_seen[a] = next_aff
        next_aff += 1

      logging.debug('added %i to %s', affs_seen[a], n)
      n_affs.add(affs_seen[a])

    if 'Additional' in row and len(row['Additional']) > 0:
      code, value = row['Additional'].split('=')
      additional[code] = value
      n_additional.add(code)

    # write the author
    p.add_run(n) #.italic = True
    
    # author afiliations
    aff_txt = ','.join(sorted([str(x) for x in n_affs]) + sorted([str(x) for x in n_additional]))

    run = p.add_run(aff_txt)
    #run.italic = True
    run.font.superscript = True

  # now write all the affs
  p = document.add_paragraph()

  for i in range(1, next_aff):
    if i > 1:
      p.add_run('; ') #.italic = True

    num = p.add_run(str(i))
    num.font.superscript = True
    #num.italic = True

    text = affs[i]
    p.add_run(text) #.italic = True

  # additional
  if len(additional) > 0:
    p = document.add_paragraph()
    p.add_run('; '.join(['{} {}'.format(x, additional[x]) for x in additional]))

  document.save(output)

  logging.info('any similar affiliations?')
  for x in sorted(affs_seen.keys()):
    sys.stdout.write('{}: {}\n'.format(x, affs_seen[x]))

  logging.info('done')

if __name__ == '__main__':
  parser = argparse.ArgumentParser(description='Assess MSI')
  parser.add_argument('--input', required=True, help='filename')
  parser.add_argument('--output', required=False, default='data.docx', help='filename')
  parser.add_argument('--name', required=True, help='name field')
  parser.add_argument('--title', required=True, help='title')
  parser.add_argument('--affiliations', nargs='+', required=True, help='affiliations fields')
  parser.add_argument('--verbose', action='store_true', help='more logging')
  args = parser.parse_args()
  if args.verbose:
    logging.basicConfig(format='%(asctime)s %(levelname)s %(message)s', level=logging.DEBUG)
  else:
    logging.basicConfig(format='%(asctime)s %(levelname)s %(message)s', level=logging.INFO)

  main(args.input, args.name, args.affiliations, args.title, args.output)
